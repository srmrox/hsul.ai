#!/usr/bin/env python3
"""
Policy Manual Content Generation
This script generates content for sections in an initialized policy manual project.
Stage 2 of the multi-stage policy manual generation workflow.
"""

import requests
import json
import time
import os
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


@dataclass
class SectionInfo:
    """Data class to hold section information"""
    number: str
    title: str
    description: str = ""
    content: str = ""
    status: str = "pending"
    word_count: int = 0
    review_notes: str = ""
    needs_revision: bool = False

    def to_dict(self):
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data):
        return cls(
            number=data.get('number', ''),
            title=data.get('title', ''),
            description=data.get('description', ''),
            content=data.get('content', ''),
            status=data.get('status', 'pending'),
            word_count=data.get('word_count', 0),
            review_notes=data.get('review_notes', ''),
            needs_revision=data.get('needs_revision', False)
        )


class LMStudioClient:
    """Client for connecting to LM Studio API"""
    
    def __init__(self, base_url: str = "http://localhost:1234", model_name: str = "local-model"):
        self.base_url = base_url.rstrip('/')
        self.model_name = model_name
        self.api_url = f"{self.base_url}/v1/chat/completions"
        
    def test_connection(self) -> bool:
        """Test if the connection to LM Studio is working"""
        try:
            response = requests.get(f"{self.base_url}/v1/models", timeout=10)
            return response.status_code == 200
        except requests.RequestException as e:
            print(f"Connection test failed: {e}")
            return False
    
    def generate_response(self, prompt: str, max_tokens: int = 2000, temperature: float = 0.7) -> Optional[str]:
        """Generate a response from LM Studio"""
        payload = {
            "model": self.model_name,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False
        }
        
        try:
            response = requests.post(self.api_url, json=payload, timeout=120)
            response.raise_for_status()
            result = response.json()
            return result['choices'][0]['message']['content'].strip()
        except requests.RequestException as e:
            print(f"API request failed: {e}")
            return None
        except (KeyError, IndexError) as e:
            print(f"Error parsing response: {e}")
            return None


class DocumentGenerator:
    """Generates RTF and DOCX formatted documents"""
    
    def create_rtf_manual(self, sections: List[SectionInfo], 
                         manual_description: str, variables: Dict,
                         filename: str = "policy_manual.rtf") -> str:
        """Create a manually formatted RTF policy manual"""
        rtf_content = []
        
        # RTF Header
        rtf_content.append(r"{\rtf1\ansi\deff0")
        rtf_content.append(r"{\fonttbl{\f0 Times New Roman;}\f1 Arial;}")
        rtf_content.append(r"{\colortbl;\red0\green0\blue0;}")
        
        # Title
        rtf_content.append(r"\pard\qc\f0\fs36\b POLICY MANUAL\b0\fs24\par")
        rtf_content.append(r"\par")
        
        # Description
        rtf_content.append(rf"\pard\qc\f0\fs20\b Subject: {manual_description}\b0\par")
        rtf_content.append(r"\par")
        
        # Date
        date_str = datetime.now().strftime('%B %d, %Y')
        rtf_content.append(rf"\pard\qc\f0\fs16 Generated: {date_str}\par")
        rtf_content.append(r"\par\par")
        
        # Table of Contents
        rtf_content.append(r"\pard\ql\f0\fs24\b\ul TABLE OF CONTENTS\ul0\b0\par")
        rtf_content.append(r"\par")
        
        for section in sections:
            indent = "\\tab " * (section.number.count('.') - 1)
            rtf_content.append(rf"\pard\ql\f0\fs18{indent}{section.number}. {section.title}\par")
        
        rtf_content.append(r"\par\par")
        rtf_content.append(r"\page")  # Page break
        
        # Add sections
        for section in sections:
            if section.content.strip():
                # Section header
                section_header = f"Section {section.number}: {section.title}"
                rtf_content.append(rf"\pard\ql\f0\fs20\b\ul {section_header}\ul0\b0\par")
                rtf_content.append(r"\par")
                
                # Section content - escape RTF special characters and substitute variables
                content = section.content
                # Substitute variables
                for var_name, var_info in variables.items():
                    placeholder = f"[{var_name.upper().replace('_', ' ')}]"
                    if placeholder in content:
                        content = content.replace(placeholder, var_info.get('default_value', placeholder))
                
                content = content.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                
                # Split content into paragraphs
                paragraphs = content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        clean_para = paragraph.replace('\n', ' ').strip()
                        rtf_content.append(rf"\pard\ql\f0\fs14 {clean_para}\par")
                        rtf_content.append(r"\par")
                
                rtf_content.append(r"\par")
        
        # RTF footer
        rtf_content.append("}")
        
        # Write RTF file
        with open(filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(rtf_content))
        
        return filename
    
    def create_docx_manual(self, sections: List[SectionInfo], 
                          manual_description: str, variables: Dict,
                          filename: str = "policy_manual.docx") -> str:
        """Create a DOCX formatted policy manual"""
        doc = Document()
        
        # Title
        title = doc.add_heading('POLICY MANUAL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Description
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(f'Subject: {manual_description}')
        desc_run.bold = True
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Table of Contents
        toc_heading = doc.add_heading('TABLE OF CONTENTS', level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for section in sections:
            toc_para = doc.add_paragraph()
            indent_level = section.number.count('.') - 1
            toc_para.paragraph_format.left_indent = Inches(0.5 * indent_level)
            toc_para.add_run(f'{section.number}. {section.title}')
        
        doc.add_page_break()
        
        # Add sections
        for section in sections:
            if section.content.strip():
                # Section header
                section_header = f'Section {section.number}: {section.title}'
                heading = doc.add_heading(section_header, level=2)
                
                # Section content with variable substitution
                content = section.content
                for var_name, var_info in variables.items():
                    placeholder = f"[{var_name.upper().replace('_', ' ')}]"
                    if placeholder in content:
                        content = content.replace(placeholder, var_info.get('default_value', placeholder))
                
                # Split content into paragraphs
                paragraphs = content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        clean_para = paragraph.replace('\n', ' ').strip()
                        para = doc.add_paragraph(clean_para)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Save document
        doc.save(filename)
        return filename


class ContentGenerator:
    """Main class for generating policy manual content"""
    
    def __init__(self, project_dir: str):
        self.project_dir = project_dir
        # Use common config.json in main directory, project-specific files in project directory
        self.config_file = os.path.join(os.getcwd(), "config.json")
        # Use common variables.json in main directory
        self.variables_file = os.path.join(os.getcwd(), "variables.json")
        # Use common organogram.json in main directory  
        self.organogram_file = os.path.join(os.getcwd(), "organogram.json")
        self.sections_file = os.path.join(project_dir, "sections.json")
        self.status_file = os.path.join(project_dir, "status.json")
        
        self.sections: List[SectionInfo] = []
        self.config = {}
        self.variables = {}
        self.organogram = {}
        self.client = None
    
    def load_project(self):
        """Load project configuration and data"""
        # Load common config (LM Studio settings)
        with open(self.config_file, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
        
        # Load project-specific data from status.json and merge into config
        if os.path.exists(self.status_file):
            with open(self.status_file, 'r', encoding='utf-8') as f:
                status_data = json.load(f)
                # Merge project-specific fields into config for backward compatibility
                self.config.update({
                    'manual_description': status_data.get('manual_description', ''),
                    'created_date': status_data.get('created_date', ''),
                    'project_name': status_data.get('project_name', ''),
                    'stage': status_data.get('stage', ''),
                    'policy_type': status_data.get('policy_type', ''),
                    'responsibilities': status_data.get('responsibilities', {})
                })
        
        # Load sections
        with open(self.sections_file, 'r', encoding='utf-8') as f:
            sections_data = json.load(f)
            self.sections = [SectionInfo.from_dict(data) for data in sections_data]
        
        # Load variables
        with open(self.variables_file, 'r', encoding='utf-8') as f:
            variables_data = json.load(f)
            self.variables = variables_data
        
        # Load organogram
        if os.path.exists(self.organogram_file):
            with open(self.organogram_file, 'r', encoding='utf-8') as f:
                self.organogram = json.load(f)
        
        # Initialize LM Studio client
        self.client = LMStudioClient(
            self.config.get('lm_studio_url', 'http://localhost:1234'),
            self.config.get('model_name', 'local-model')
        )
    
    def save_sections(self):
        """Save sections to file"""
        sections_data = [section.to_dict() for section in self.sections]
        with open(self.sections_file, 'w', encoding='utf-8') as f:
            json.dump(sections_data, f, indent=2, ensure_ascii=False)
    
    def save_status(self, status: Dict):
        """Save current project status"""
        status['last_updated'] = datetime.now().isoformat()
        with open(self.status_file, 'w', encoding='utf-8') as f:
            json.dump(status, f, indent=2)
    
    def generate_section_content(self, section_index: int, existing_content: str = "") -> str:
        """Generate content for a specific section using description as guidance"""
        if section_index >= len(self.sections):
            return ""
            
        section = self.sections[section_index]
        manual_description = self.config.get('manual_description', '')
        
        context_prompt = ""
        if existing_content:
            # Limit context to avoid token limits
            context_words = existing_content.split()
            if len(context_words) > 500:
                context_prompt = f"\n\nPreviously written content (for context):\n{' '.join(context_words[-500:])}"
            else:
                context_prompt = f"\n\nPreviously written content (for context):\n{existing_content}"
        
        # Prepare variables context
        variables_context = ""
        if self.variables:
            variables_context = "\n\nAvailable variables (use these placeholders in your content):\n"
            for var_name, var_info in self.variables.items():
                placeholder = f"[{var_name.upper().replace('_', ' ')}]"
                variables_context += f"- {placeholder}: {var_info.get('description', 'No description')}\n"
            variables_context += "\nUse these placeholders where appropriate in your content (e.g., [COMPANY NAME], [EFFECTIVE DATE])."
        
        # Prepare organogram/responsibility context
        responsibility_context = ""
        responsibilities = self.config.get('responsibilities', {})
        if responsibilities:
            responsibility_context = "\n\nKey roles and responsibilities for this policy:\n"
            for role_type, role_info in responsibilities.items():
                responsibility_context += f"- {role_info.get('title', role_type)}: {role_info.get('name', '[NAME]')} ({role_info.get('email', '[EMAIL]')})\n"
            responsibility_context += "\nReference these roles when specifying responsibilities, approvals, or escalation procedures in your content."
        
        prompt = f"""
        You are writing a comprehensive policy manual about: {manual_description}
        
        Please write detailed content for this section:
        Section {section.number}: {section.title}
        
        Section Description (your guide for what to cover):
        {section.description}
        {variables_context}
        {responsibility_context}
        
        Requirements:
        - Write 400-1000 words of detailed, professional content
        - Use clear, policy-appropriate language suitable for a professional environment
        - Include specific guidelines, procedures, or requirements as relevant
        - Structure with subheadings if appropriate based on the section description
        - Be comprehensive and address all aspects mentioned in the description
        - Do not repeat information from other sections
        - Use professional policy language with clear directives and actionable guidance
        - Include relevant examples or scenarios where appropriate
        - Use variable placeholders where appropriate (e.g., [COMPANY NAME], [EFFECTIVE DATE])
        - Reference appropriate roles for responsibilities and approvals when relevant
        {context_prompt}
        
        Write only the section content, do not include the section number or title in your response.
        Focus specifically on what is described in the section description above.
        """
        
        print(f"🔄 Generating content for Section {section.number}: {section.title}")
        print(f"   📝 Using description: {section.description[:100]}...")
        
        content = self.client.generate_response(prompt, max_tokens=1500, temperature=0.6)
        
        if content:
            # Update section
            self.sections[section_index].content = content
            self.sections[section_index].status = 'generated'
            self.sections[section_index].word_count = len(content.split())
            
            print(f"✅ Generated {len(content.split())} words for section {section.number}")
        else:
            print(f"❌ Failed to generate content for section {section.number}")
        
        return content or ""
    
    def generate_all_sections(self, resume_from: int = 0):
        """Generate content for all sections starting from a specific index"""
        print(f"\n🔄 Starting content generation from section {resume_from + 1} of {len(self.sections)}...")
        
        existing_content = ""
        
        # Build existing content from completed sections
        for i in range(resume_from):
            if i < len(self.sections) and self.sections[i].content:
                existing_content += f"\n\nSection {self.sections[i].number}: {self.sections[i].title}\n{self.sections[i].content}"
        
        for i in range(resume_from, len(self.sections)):
            section_content = self.generate_section_content(i, existing_content)
            
            if section_content:
                existing_content += f"\n\nSection {self.sections[i].number}: {self.sections[i].title}\n{section_content}"
            
            # Save progress after each section
            self.save_sections()
            
            completed_count = sum(1 for s in self.sections if s.status == 'generated')
            self.save_status({
                'phase': 'generating_content',
                'total_sections': len(self.sections),
                'completed_sections': completed_count,
                'current_section': i,
                'manual_description': self.config.get('manual_description', '')
            })
            
            # Small delay to avoid overwhelming the API
            time.sleep(1)
        
        print("✅ Completed content generation for all sections")
    
    def save_to_files(self, base_filename: str):
        """Save the manual to RTF, DOCX, and JSON formats"""
        # Save sections data to JSON
        json_filename = f"{base_filename}_data.json"
        sections_data = [section.to_dict() for section in self.sections]
        with open(json_filename, 'w', encoding='utf-8') as f:
            json.dump({
                'manual_description': self.config.get('manual_description', ''),
                'generated_date': datetime.now().isoformat(),
                'sections': sections_data,
                'variables': self.variables,
                'statistics': {
                    'total_sections': len(self.sections),
                    'total_words': sum(s.word_count for s in self.sections),
                    'sections_needing_revision': sum(1 for s in self.sections if s.needs_revision)
                }
            }, f, indent=2, ensure_ascii=False)
        print(f"📁 Saved data to {json_filename}")
        
        # Save complete manual to RTF and DOCX files
        doc_generator = DocumentGenerator()
        
        rtf_filename = f"{base_filename}.rtf"
        doc_generator.create_rtf_manual(self.sections, 
                                      self.config.get('manual_description', ''), 
                                      self.variables, rtf_filename)
        print(f"📁 Saved RTF manual to {rtf_filename}")
        
        docx_filename = f"{base_filename}.docx"
        doc_generator.create_docx_manual(self.sections, 
                                       self.config.get('manual_description', ''), 
                                       self.variables, docx_filename)
        print(f"📁 Saved DOCX manual to {docx_filename}")
        
        return json_filename, rtf_filename, docx_filename


def list_available_projects() -> List[str]:
    """List available project directories"""
    projects = []
    for item in os.listdir('.'):
        if os.path.isdir(item) and item.endswith('_project'):
            project_name = item.replace('_project', '')
            status_file = os.path.join(item, 'status.json')
            sections_file = os.path.join(item, 'sections.json')
            
            if os.path.exists(status_file) and os.path.exists(sections_file):
                projects.append(project_name)
    
    return projects


def main():
    """Main function for content generation"""
    print("🚀 Policy Manual Content Generation")
    print("=" * 50)
    
    # List available projects
    projects = list_available_projects()
    
    if not projects:
        print("❌ No initialized projects found!")
        print("Run 'python init_project.py' first to create a project.")
        return
    
    print("📁 Available projects:")
    for i, project in enumerate(projects, 1):
        project_dir = f"{project}_project"
        try:
            with open(os.path.join(project_dir, 'status.json'), 'r') as f:
                status = json.load(f)
                description = status.get('manual_description', 'No description')
                created = status.get('created_date', 'Unknown')
                phase = status.get('phase', 'unknown')
                completed = status.get('completed_sections', 0)
                total = status.get('total_sections', 0)
            
            print(f"   {i}. {project}")
            print(f"      Description: {description}")
            print(f"      Status: {phase} ({completed}/{total} sections)")
            print(f"      Created: {created[:10] if created != 'Unknown' else 'Unknown'}")
            print()
            
        except Exception as e:
            print(f"   {i}. {project} (Error loading info: {e})")
    
    # Get user selection
    while True:
        try:
            choice = input(f"Select project (1-{len(projects)}): ").strip()
            project_index = int(choice) - 1
            
            if 0 <= project_index < len(projects):
                selected_project = projects[project_index]
                break
            else:
                print("Invalid selection!")
        except ValueError:
            print("Please enter a valid number!")
    
    project_dir = f"{selected_project}_project"
    print(f"\n📂 Selected project: {selected_project}")
    
    # Initialize content generator
    generator = ContentGenerator(project_dir)
    
    try:
        # Load project data
        print("📋 Loading project data...")
        generator.load_project()
        
        # Test LM Studio connection
        print("🔌 Testing LM Studio connection...")
        if not generator.client.test_connection():
            print("❌ Failed to connect to LM Studio. Please check the URL and ensure LM Studio is running.")
            return
        
        print("✅ Successfully connected to LM Studio!")
        
        # Determine where to start/resume
        completed_sections = sum(1 for s in generator.sections if s.status == 'generated')
        
        if completed_sections > 0:
            print(f"📊 Found {completed_sections}/{len(generator.sections)} completed sections")
            resume = input("Resume from where you left off? (y/n): ").strip().lower()
            
            if resume == 'y':
                start_from = completed_sections
                print(f"🔄 Resuming from section {start_from + 1}")
            else:
                start_from = 0
                print("🔄 Starting from the beginning")
        else:
            start_from = 0
            print("🔄 Starting content generation")
        
        # Generate content
        generator.generate_all_sections(start_from)
        
        # Save final results
        print("\n💾 Saving final results...")
        json_file, rtf_file, docx_file = generator.save_to_files(selected_project)
        
        # Update final status
        generator.save_status({
            'phase': 'content_generated',
            'total_sections': len(generator.sections),
            'completed_sections': len(generator.sections),
            'manual_description': generator.config.get('manual_description', ''),
            'generation_completed': datetime.now().isoformat()
        })
        
        print(f"\n🎉 Content generation completed!")
        print(f"📄 Files created:")
        print(f"   - {json_file} (structured data)")
        print(f"   - {rtf_file} (RTF formatted manual)")
        print(f"   - {docx_file} (Word document)")
        
        # Show final statistics
        total_words = sum(s.word_count for s in generator.sections)
        print(f"\n📊 Final Statistics:")
        print(f"   Total sections: {len(generator.sections)}")
        print(f"   Total words: {total_words:,}")
        print(f"   Variables used: {len(generator.variables)}")
        
    except Exception as e:
        print(f"❌ Error during content generation: {e}")
        return


if __name__ == "__main__":
    main()