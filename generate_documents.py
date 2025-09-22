#!/usr/bin/env python3
"""
Simple document generator for the GUI system.
"""

import json
import os
import argparse
from datetime import datetime
from config_manager import get_config_manager

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def main():
    parser = argparse.ArgumentParser(description='Generate documents')
    parser.add_argument('--output', type=str, help='Output directory')
    args = parser.parse_args()
    
    print(" Stage 4: Document Generation")
    
    if not DOCX_AVAILABLE:
        print(" python-docx required")
        return
    
    # Initialize configuration manager for variable substitution
    config_manager = get_config_manager()
    
    output_dir = args.output if args.output else "generated_documents"
    os.makedirs(output_dir, exist_ok=True)
    
    content_files = [f for f in os.listdir('.') if f.startswith('content_') and f.endswith('.json')]
    
    if not content_files:
        print(" No content files found")
        return
    
    for content_file in content_files:
        try:
            with open(content_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            doc = Document()
            title = data.get('manual_description', 'Policy Manual')
            # Apply variable substitution to title
            title = config_manager.apply_templates(title)
            doc.add_heading(title, 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            
            sections = data.get('sections', {})
            for key, section in sections.items():
                if isinstance(section, dict) and 'content' in section:
                    section_title = section.get('title', key)
                    # Apply variable substitution to section title
                    section_title = config_manager.apply_templates(section_title)
                    doc.add_heading(section_title, level=1)
                    
                    content = section.get('content', '')
                    if content:
                        # Apply variable substitution to content
                        content = config_manager.apply_templates(content)
                        doc.add_paragraph(content)
            
            base_name = content_file.replace('content_', '').replace('.json', '')
            output_file = os.path.join(output_dir, f"{base_name}_manual.docx")
            doc.save(output_file)
            print(f" Created: {os.path.basename(output_file)}")
            
        except Exception as e:
            print(f" Error: {e}")
    
    print(" Complete!")


if __name__ == "__main__":
    main()
